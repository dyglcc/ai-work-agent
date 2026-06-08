"""文档解析工具：提取 PDF、DOCX、TXT、CSV、Markdown 等格式的文本内容"""
from __future__ import annotations

import io
import logging
import pathlib
from typing import Any

logger = logging.getLogger(__name__)

# 支持的文档类型
SUPPORTED_EXTENSIONS = {
    ".txt", ".csv", ".md", ".markdown", ".json", ".xml", ".html", ".htm",
    ".py", ".js", ".ts", ".java", ".go", ".rs", ".c", ".cpp", ".h",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf",
    ".docx", ".pdf", ".rtf",
}


def parse_text_file(content: bytes, encoding: str = "utf-8") -> str:
    """解析纯文本文件，自动检测编码"""
    try:
        return content.decode(encoding, errors="replace")
    except (UnicodeDecodeError, LookupError):
        # 尝试常见编码
        for enc in ["gbk", "gb2312", "latin-1", "cp1252"]:
            try:
                return content.decode(enc, errors="replace")
            except (UnicodeDecodeError, LookupError):
                continue
        return content.decode("utf-8", errors="replace")


def parse_docx(content: bytes) -> str:
    """解析 DOCX 文件，提取所有段落和表格文本"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        parts = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 保留标题层级
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    level = para.style.name.split()[-1]
                    parts.append(f"{'#' * int(level)} {text}")
                else:
                    parts.append(text)

        # 提取表格
        for table in doc.tables:
            parts.append("")  # 空行分隔
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace("\n", " ")
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n".join(parts)
    except Exception as e:
        logger.exception("解析 DOCX 文件失败")
        raise ValueError(f"DOCX 解析失败: {e}")


def parse_pdf(content: bytes) -> str:
    """解析 PDF 文件，提取文本内容"""
    try:
        # 尝试 PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    parts.append(text.strip())
            if parts:
                return "\n\n".join(parts)
        except ImportError:
            pass

        # 尝试 pdfplumber (更准确的文本提取)
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        parts.append(text.strip())
            if parts:
                return "\n\n".join(parts)
        except ImportError:
            pass

        raise ImportError("未安装 PDF 解析库，请安装 PyPDF2 或 pdfplumber")
    except ImportError as e:
        raise ImportError(str(e))
    except Exception as e:
        logger.exception("解析 PDF 文件失败")
        raise ValueError(f"PDF 解析失败: {e}")


def parse_csv(content: bytes) -> list[dict[str, str]]:
    """解析 CSV 文件，返回字典列表"""
    import csv

    text = parse_text_file(content)
    reader = csv.DictReader(io.StringIO(text))
    return list(reader)


def parse_markdown(content: bytes) -> dict[str, Any]:
    """解析 Markdown 文件，提取标题、段落和代码块"""
    text = parse_text_file(content)
    lines = text.split("\n")

    result = {
        "title": "",
        "headings": [],
        "paragraphs": [],
        "code_blocks": [],
        "plain_text": text,
    }

    in_code_block = False
    code_lang = ""
    code_lines = []
    current_paragraph = []

    for line in lines:
        # 标题
        if line.startswith("# ") and not result["title"]:
            result["title"] = line[2:].strip()
            result["headings"].append({"level": 1, "text": line[2:].strip()})
            continue
        elif line.startswith("## ") or line.startswith("### "):
            level = 2 if line.startswith("## ") else 3
            result["headings"].append({"level": level, "text": line[level + 1:].strip()})
            if current_paragraph:
                result["paragraphs"].append(" ".join(current_paragraph))
                current_paragraph = []
            continue

        # 代码块
        if line.startswith("```"):
            if in_code_block:
                result["code_blocks"].append({
                    "language": code_lang,
                    "code": "\n".join(code_lines),
                })
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lang = line[3:].strip()
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # 普通段落
        stripped = line.strip()
        if stripped:
            current_paragraph.append(stripped)
        elif current_paragraph:
            result["paragraphs"].append(" ".join(current_paragraph))
            current_paragraph = []

    if current_paragraph:
        result["paragraphs"].append(" ".join(current_paragraph))

    return result


def parse_document(
    content: bytes,
    filename: str,
    max_chars: int = 10000,
) -> dict[str, Any]:
    """
    统一的文档解析入口，根据文件扩展名自动选择解析器。

    Args:
        content: 文件二进制内容
        filename: 文件名（含扩展名）
        max_chars: 最大返回字符数

    Returns:
        {
            "success": True/False,
            "filename": str,
            "file_type": str,
            "text": str,           # 提取的文本内容
            "char_count": int,     # 文本字符数
            "truncated": bool,     # 是否被截断
            "metadata": dict,      # 文件元数据
            "error": str | None,
        }
    """
    ext = pathlib.Path(filename).suffix.lower()
    text = ""
    metadata: dict[str, Any] = {}

    if not is_supported(filename):
        return {
            "success": False,
            "filename": filename,
            "file_type": ext,
            "text": "",
            "char_count": 0,
            "truncated": False,
            "metadata": {},
            "error": f"不支持的文件类型: {ext}",
        }

    try:
        if ext == ".docx":
            text = parse_docx(content)
            metadata["parser"] = "python-docx"
        elif ext == ".pdf":
            text = parse_pdf(content)
            metadata["parser"] = "PyPDF2/pdfplumber"
        elif ext == ".csv":
            rows = parse_csv(content)
            text = "\n".join(
                ", ".join(f"{k}: {v}" for k, v in row.items())
                for row in rows
            )
            metadata["parser"] = "csv"
            metadata["row_count"] = len(rows)
        elif ext in (".md", ".markdown"):
            md_result = parse_markdown(content)
            text = md_result["plain_text"]
            metadata["parser"] = "markdown"
            metadata["title"] = md_result["title"]
            metadata["heading_count"] = len(md_result["headings"])
            metadata["code_block_count"] = len(md_result["code_blocks"])
        elif ext in (".txt", ".json", ".xml", ".html", ".htm", ".yaml", ".yml",
                     ".toml", ".ini", ".cfg", ".conf",
                     ".py", ".js", ".ts", ".java", ".go", ".rs", ".c", ".cpp", ".h"):
            text = parse_text_file(content)
            metadata["parser"] = "text"
        else:
            # 尝试作为纯文本解析
            text = parse_text_file(content)
            if len(text) < 10 and any(b < 0x20 and b not in (0x09, 0x0a, 0x0d) for b in content[:100]):
                return {
                    "success": False,
                    "filename": filename,
                    "file_type": ext,
                    "text": "",
                    "char_count": 0,
                    "truncated": False,
                    "metadata": {},
                    "error": f"不支持的文件类型: {ext}（可能是二进制文件）",
                }
            metadata["parser"] = "text (fallback)"

        # 截断
        truncated = False
        if len(text) > max_chars:
            text = text[:max_chars]
            truncated = True

        return {
            "success": True,
            "filename": filename,
            "file_type": ext,
            "text": text,
            "char_count": len(text),
            "truncated": truncated,
            "metadata": metadata,
            "error": None,
        }

    except ImportError as e:
        return {
            "success": False,
            "filename": filename,
            "file_type": ext,
            "text": "",
            "char_count": 0,
            "truncated": False,
            "metadata": {},
            "error": f"缺少解析库: {e}",
        }
    except Exception as e:
        logger.exception("解析文档 %s 失败", filename)
        return {
            "success": False,
            "filename": filename,
            "file_type": ext,
            "text": "",
            "char_count": 0,
            "truncated": False,
            "metadata": {},
            "error": str(e),
        }


def is_supported(filename: str) -> bool:
    """检查文件类型是否受支持"""
    ext = pathlib.Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS


def get_supported_types() -> list[str]:
    """返回支持的文件扩展名列表"""
    return sorted(SUPPORTED_EXTENSIONS)