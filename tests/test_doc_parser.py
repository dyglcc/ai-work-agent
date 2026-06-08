"""文档解析模块测试"""

import pytest
from app.services.doc_parser import (
    parse_document,
    parse_text_file,
    parse_docx,
    parse_pdf,
    parse_markdown,
    parse_csv,
    is_supported,
    get_supported_types,
)


class TestIsSupported:
    """测试文件类型支持检查"""

    def test_supported_extensions(self):
        supported = get_supported_types()
        assert ".txt" in supported
        assert ".docx" in supported
        assert ".pdf" in supported
        assert ".md" in supported
        assert ".html" in supported
        assert ".htm" in supported
        assert ".csv" in supported
        assert ".json" in supported
        assert ".py" in supported
        assert ".js" in supported
        assert ".ts" in supported

    def test_unsupported_extension(self):
        assert not is_supported("file.xyz")
        assert not is_supported("file.abc")
        assert not is_supported("file")

    def test_supported_filename(self):
        assert is_supported("document.txt")
        assert is_supported("report.docx")
        assert is_supported("slides.pdf")
        assert is_supported("README.md")
        assert is_supported("index.html")
        assert is_supported("data.csv")
        assert is_supported("config.json")
        assert is_supported("main.py")
        assert is_supported("app.js")
        assert is_supported("types.ts")
        assert is_supported("page.htm")

    def test_case_insensitive(self):
        assert is_supported("FILE.TXT")
        assert is_supported("File.Docx")
        assert is_supported("DOCUMENT.PDF")


class TestParseTextFile:
    """测试纯文本解析"""

    def test_parse_utf8_text(self):
        content = "Hello World 你好世界".encode("utf-8")
        result = parse_text_file(content)
        assert result == "Hello World 你好世界"

    def test_parse_python(self):
        content = b"def hello():\n    print('Hello')\n"
        result = parse_text_file(content)
        assert "def hello" in result

    def test_parse_empty_file(self):
        content = b""
        result = parse_text_file(content)
        assert result == ""

    def test_parse_large_file(self):
        content = ("A" * 100000).encode("utf-8")
        result = parse_text_file(content)
        assert len(result) == 100000


class TestParseMarkdown:
    """测试 Markdown 解析"""

    def test_parse_markdown(self):
        content = b"# Title\n\nParagraph text\n\n- Item 1\n- Item 2"
        result = parse_markdown(content)
        assert result["title"] == "Title"
        assert "Paragraph text" in result["plain_text"]
        assert "Item 1" in result["plain_text"]

    def test_parse_markdown_with_code(self):
        content = b"# Title\n\n```python\nprint('hello')\n```\n\nEnd"
        result = parse_markdown(content)
        assert "print('hello')" in result["plain_text"]
        assert len(result["code_blocks"]) >= 1
        assert result["code_blocks"][0]["code"] == "print('hello')"


class TestParseDocx:
    """测试 DOCX 解析"""

    def test_parse_docx(self):
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("Another paragraph with content.")

        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        result = parse_docx(content)
        assert "Test Document" in result
        assert "test paragraph" in result

    def test_parse_docx_empty(self):
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument()
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        result = parse_docx(content)
        assert isinstance(result, str)

    def test_parse_docx_invalid(self):
        content = b"not a valid docx file"
        with pytest.raises(ValueError, match="DOCX 解析失败"):
            parse_docx(content)


class TestParsePDF:
    """测试 PDF 解析"""

    def test_parse_pdf_invalid(self):
        content = b"not a valid pdf"
        with pytest.raises((ValueError, ImportError)):
            parse_pdf(content)


class TestParseCSV:
    """测试 CSV 解析"""

    def test_parse_csv(self):
        content = b"name,age,city\nAlice,30,Beijing\nBob,25,Shanghai"
        result = parse_csv(content)
        assert len(result) == 2
        assert result[0]["name"] == "Alice"
        assert result[0]["age"] == "30"
        assert result[1]["name"] == "Bob"


class TestParseDocument:
    """测试统一解析入口"""

    def test_parse_txt(self):
        content = "Hello World".encode("utf-8")
        result = parse_document(content, "test.txt")
        assert result["success"] is True
        assert result["text"] == "Hello World"
        assert result["filename"] == "test.txt"
        assert "metadata" in result

    def test_parse_md(self):
        content = b"# Title\n\nContent"
        result = parse_document(content, "README.md")
        assert result["success"] is True
        assert "Title" in result["text"]

    def test_parse_html(self):
        content = b"<html><body><p>Test</p></body></html>"
        result = parse_document(content, "page.html")
        assert result["success"] is True

    def test_parse_unsupported(self):
        content = b"some binary data"
        result = parse_document(content, "file.xyz")
        assert result["success"] is False
        assert "不支持" in result.get("error", "")

    def test_parse_python_file(self):
        content = b"import os\n\nprint('hello')\n"
        result = parse_document(content, "script.py")
        assert result["success"] is True
        assert "import os" in result["text"]

    def test_parse_js_file(self):
        content = b"const x = 1;\nconsole.log(x);\n"
        result = parse_document(content, "app.js")
        assert result["success"] is True
        assert "const x = 1" in result["text"]
